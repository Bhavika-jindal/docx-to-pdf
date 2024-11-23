from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from reportlab.pdfgen import canvas
import os


app = Flask(__name__)

# CORS Configuration to allow frontend to interact with the backend
CORS(app, resources={r"/upload": {"origins": "http://localhost:3000"}})


# Ensure the uploads folder exists
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER



# Increase the maximum content length to 16MB for file uploads
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

@app.route('/', methods=['GET'])
def home():
    return "Welcome to the Docx-to-PDF Converter API! Use the /upload route to upload a file."

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/upload', methods=['POST'])
def upload_file():
    print("Request files:", request.files)  # Debugging to print the files sent in the request

    if 'file' not in request.files:
        print("No file found in the request")
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    print(f"File received: {file.filename}")  # Confirm the file is received
    if file.filename == '' or not file.filename.endswith('.docx'):
        print("Invalid file format or empty filename")
        return jsonify({"error": "Invalid file format"}), 400

    temp_doc_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(temp_doc_path)
    print(f"File saved at: {temp_doc_path}")

    metadata = extract_metadata(temp_doc_path)
    pdf_path = convert_to_pdf(temp_doc_path)

    return jsonify({
        "message": "File uploaded successfully",
        "metadata": metadata,
        "pdf_path": pdf_path
    })



def extract_metadata(file_path):
    doc = Document(file_path)
    metadata = {
        "author": doc.core_properties.author,
        "title": doc.core_properties.title,
        "word_count": len(doc.paragraphs)
    }
    return metadata

def convert_to_pdf(file_path):
    output_pdf_path = os.path.splitext(file_path)[0] + ".pdf"
    doc = Document(file_path)
    c = canvas.Canvas(output_pdf_path)
    y = 800

    for paragraph in doc.paragraphs:
        c.drawString(100, y, paragraph.text)
        y -= 20

    c.save()
    return output_pdf_path

@app.route('/download/<filename>', methods=['GET'])
def download_pdf(filename):
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(pdf_path):
        return jsonify({"error": "File not found"}), 404

    return send_file(pdf_path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)

